from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

def heading(text, level=1):
    doc.add_heading(text, level=level)

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

heading("MongoDB Inventory Audit – Human–Computer Interaction (HCI)", 0)
p = para(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

heading("Overview", 1)
para("This document presents the HCI of the MongoDB Inventory Audit tool. "
     "The HCI is implemented as a Tkinter GUI in `a.py` (launched with `python a.py --gui`) "
     "and a richer GUI in `withgui.py`. It enables non-technical users to run audits without the CLI.")

heading("What is HCI in this Context", 1)
para("HCI refers to the interface elements and workflows that let a human perform audit tasks via interactive controls "
     "rather than command-line input. The GUI provides clear controls, feedback, and error handling.")

heading("User Workflow", 1)
for s in [
    "Connect: App shows MongoDB connection status and available collections.",
    "Configure: Select collection, set Sample Size and Threshold ($).",
    "Run Audit: Execute price testing; results appear in a table.",
    "Export: Save results to CSV.",
    "Cleanup: Optionally remove `extendedValue` fields from documents."
]:
    bullet(s)

heading("Key UI Components", 1)
for s in [
    "Collection dropdown",
    "Sample Size spinbox (1–500)",
    "Threshold ($) spinbox",
    "Run Audit / Export CSV / Cleanup Extended Values buttons",
    "Results table with Item ID, Description, Price, Qty, Value",
    "Summary line showing item count and total sampled value"
]:
    bullet(s)

heading("Design Rationale", 1)
para("The GUI improves usability, reduces input errors, and standardizes audit steps. "
     "It supports consistent sampling and export for audit evidence.")

heading("Screenshots", 1)
para("Replace the image paths below with your screenshots.")
for caption, path in [
    ("Config Screen (Select collection, parameters)", "C:\\path\\to\\screenshot_config.png"),
    ("Results Table (Audit output)", "C:\\path\\to\\screenshot_results.png"),
    ("Export Confirmation", "C:\\path\\to\\screenshot_export.png"),
]:
    para(caption, bold=True)
    try:
        doc.add_picture(path, width=Inches(6.0))
    except Exception:
        para(f"(Image not found: {path})")

heading("Optional Flow Diagram (Textual)", 1)
para("User → Select Collection → Set Parameters → Run Audit → View Results → Export CSV / Cleanup")

doc.save("HCI_Report.docx")
print("HCI_Report.docx generated.")